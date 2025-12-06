"""
DOCX parser adapter using python-docx.

Migrated from: coomb-pdf/services/resume_import_service.py
"""

from __future__ import annotations

import io
import logging

from docx import Document

from domain.ports.outbound.resume_parser_port import (
    DocumentParserPort,
    ExtractedText,
    ParserError,
)

logger = logging.getLogger(__name__)


class DOCXParserAdapter(DocumentParserPort):
    """
    DOCX text extraction using python-docx.

    This adapter extracts text from Word documents,
    including paragraphs and tables.
    """

    SUPPORTED_EXTENSIONS = [".docx", ".doc"]

    async def extract_text(self, file_content: bytes) -> ExtractedText:
        """
        Extract text content from a DOCX file.

        Args:
            file_content: Raw bytes of the DOCX file

        Returns:
            ExtractedText with content and metadata

        Raises:
            ParserError: If extraction fails
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            # Extract paragraphs
            paragraphs: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            table_texts: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        table_texts.append(row_text)

            # Combine all text
            all_texts = paragraphs + table_texts
            full_text = "\n".join(all_texts)

            if not full_text or len(full_text.strip()) < 50:
                raise ParserError(
                    "DOCX extraction returned insufficient text. "
                    "The file may be empty or corrupted."
                )

            logger.info(
                f"DOCX extracted successfully: {len(full_text)} chars, "
                f"{len(paragraphs)} paragraphs, {len(doc.tables)} tables"
            )

            return ExtractedText(
                content=full_text,
                char_count=len(full_text),
                source_type="docx",
                metadata={
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                },
            )

        except ParserError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ParserError(f"Failed to extract text from DOCX: {e}") from e

    def supports_format(self, filename: str) -> bool:
        """Check if this parser supports the given file format."""
        return any(filename.lower().endswith(ext) for ext in self.SUPPORTED_EXTENSIONS)

